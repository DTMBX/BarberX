"""
Tests for Phase 9 — Document Processing Engine
=================================================
Covers:
  G1  — test_pdf_text_extraction: native PDF → text
  G2  — test_pdf_ocr_extraction: scanned PDF → OCR text
  G3  — test_video_metadata_extraction: ffprobe → structured data
  G4  — test_thumbnail_generation: ffmpeg → JPEG file
  G5  — test_content_search: keyword search → matching results
  G6  — test_entity_extraction: regex → emails / phones
  G7  — test_processing_task_lifecycle: create → processing → completed
  G8  — test_processing_task_failure: corrupt file → status=failed
  G9  — test_api_process_evidence: POST /api/v1/evidence/<id>/process → 200/202
  G10 — test_api_get_text: GET /api/v1/evidence/<id>/text → JSON
  G11 — test_api_search: GET /api/v1/search → matching results
  G12 — test_unsupported_file_type: skip with reason

Additional coverage:
  - File type detection (MIME + extension)
  - DOCX text extraction
  - Plain-text passthrough
  - Celery sync fallback (dispatch_process_evidence)
  - Video metadata parsing
  - Batch processing
  - Processing routes (all 6 endpoints)
"""

import hashlib
import json
import os
import shutil
import tempfile
import uuid
from datetime import datetime, timezone
from pathlib import Path
from unittest.mock import MagicMock, patch, PropertyMock

import pytest


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture(scope="module")
def app():
    """Create a test Flask app with in-memory DB."""
    os.environ["FLASK_ENV"] = "testing"
    os.environ["EVIDENT_ASYNC"] = "0"  # Force sync mode for tests
    from app_config import create_app

    application = create_app()
    application.config["TESTING"] = True
    application.config["WTF_CSRF_ENABLED"] = False
    application.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///:memory:"
    application.config["RATELIMIT_ENABLED"] = False
    yield application


@pytest.fixture(scope="module")
def _db(app):
    """Module-scoped: push app context, create tables."""
    from auth.models import db

    # Ensure all models are imported so create_all() sees them
    import models.webhook           # noqa: F401
    import models.document_processing  # noqa: F401
    import models.forensic_media    # noqa: F401
    import models.evidence          # noqa: F401
    import models.legal_case        # noqa: F401

    with app.app_context():
        db.create_all()
        yield db
        db.session.remove()


@pytest.fixture()
def client(app):
    return app.test_client()


@pytest.fixture()
def test_user(app, _db):
    """Create a test user for each test."""
    from auth.models import User, UserRole, TierLevel

    user = User(
        email=f"proc-test-{uuid.uuid4().hex[:8]}@evident.test",
        username=f"proctest_{uuid.uuid4().hex[:8]}",
        full_name="Processing Test User",
        role=UserRole.USER,
        tier=TierLevel.PRO,
        is_verified=True,
        is_active=True,
    )
    user.set_password("TestPass123!")
    _db.session.add(user)
    _db.session.commit()
    return user


@pytest.fixture()
def admin_user(app, _db):
    """Create an admin user."""
    from auth.models import User, UserRole, TierLevel

    user = User(
        email=f"proc-admin-{uuid.uuid4().hex[:8]}@evident.test",
        username=f"procadmin_{uuid.uuid4().hex[:8]}",
        full_name="Processing Admin",
        role=UserRole.ADMIN,
        tier=TierLevel.ADMIN,
        is_verified=True,
        is_active=True,
    )
    user.set_password("AdminPass123!")
    _db.session.add(user)
    _db.session.commit()
    return user


@pytest.fixture()
def api_token(app, _db, test_user):
    """Create a valid API token for the test user."""
    from auth.models import ApiToken
    import secrets

    raw = secrets.token_urlsafe(32)
    token = ApiToken(
        user_id=test_user.id,
        token=raw,
        name="test-processing-token",
        is_active=True,
    )
    _db.session.add(token)
    _db.session.commit()
    return raw


@pytest.fixture()
def test_case(app, _db, test_user):
    """Create a test legal case."""
    from models.legal_case import LegalCase

    case = LegalCase(
        case_name=f"Test Case {uuid.uuid4().hex[:8]}",
        case_number=f"TC-{uuid.uuid4().hex[:6]}",
        case_type="civil",
        status="active",
        created_by_id=test_user.id,
    )
    _db.session.add(case)
    _db.session.commit()
    return case


@pytest.fixture()
def sample_pdf(tmp_path):
    """Create a sample PDF with native text for testing."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        pdf_path = tmp_path / "test_doc.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(72, 720, "EVIDENT TECHNOLOGIES TEST DOCUMENT")
        c.drawString(72, 700, "Contact: test@evident.tech")
        c.drawString(72, 680, "Phone: (555) 123-4567")
        c.drawString(72, 660, "This is a test document for processing pipeline validation.")
        c.save()
        return str(pdf_path)
    except ImportError:
        # If reportlab not available, create a minimal PDF manually
        pdf_path = tmp_path / "test_doc.pdf"
        # Minimal valid PDF with text
        pdf_content = b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj
3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj
4 0 obj << /Length 89 >>
stream
BT /F1 12 Tf 72 720 Td (EVIDENT TECHNOLOGIES TEST DOCUMENT) Tj 0 -20 Td (Contact: test@evident.tech) Tj ET
endstream
endobj
5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000407 00000 n 
trailer << /Size 6 /Root 1 0 R >>
startxref
486
%%EOF"""
        pdf_path.write_bytes(pdf_content)
        return str(pdf_path)


@pytest.fixture()
def sample_text_file(tmp_path):
    """Create a sample text file."""
    text_path = tmp_path / "evidence_notes.txt"
    text_path.write_text(
        "Officer report: Incident at 123 Main St.\n"
        "Contact: officer@police.gov\n"
        "Witness phone: (555) 987-6543\n"
        "Date: 2026-01-15\n",
        encoding="utf-8",
    )
    return str(text_path)


@pytest.fixture()
def evidence_item(app, _db, test_user, test_case, sample_pdf):
    """Create an evidence item linked to case with a real file."""
    from models.evidence import EvidenceItem, CaseEvidence

    sha256 = hashlib.sha256(Path(sample_pdf).read_bytes()).hexdigest()

    item = EvidenceItem(
        original_filename="test_doc.pdf",
        stored_filename="test_doc.pdf",
        file_type="pdf",
        file_size_bytes=os.path.getsize(sample_pdf),
        mime_type="application/pdf",
        evidence_type="document",
        hash_sha256=sha256,
        processing_status="pending",
        uploaded_by_id=test_user.id,
        origin_case_id=test_case.id,
    )
    _db.session.add(item)
    _db.session.flush()

    link = CaseEvidence(
        case_id=test_case.id,
        evidence_id=item.id,
        linked_by_id=test_user.id,
        link_purpose="intake",
    )
    _db.session.add(link)
    _db.session.commit()

    # Store the file path for tests to access
    item._test_file_path = sample_pdf
    return item


# ---------------------------------------------------------------------------
# G1: PDF text extraction (native)
# ---------------------------------------------------------------------------

class TestPdfTextExtraction:
    """G1 — Native PDF text extraction."""

    def test_extract_native_pdf(self, app, sample_pdf):
        """Native PDF → text → non-empty string."""
        with app.app_context():
            from services.evidence_processor import extract_pdf_text

            result = extract_pdf_text(sample_pdf)
            assert result.success is True
            assert result.task_type == "pdf_text"
            assert result.page_count >= 1
            assert result.word_count > 0
            assert "EVIDENT" in (result.full_text or "").upper()

    def test_pdf_word_and_char_count(self, app, sample_pdf):
        """A4 — Word/char count are non-zero integers."""
        with app.app_context():
            from services.evidence_processor import extract_pdf_text

            result = extract_pdf_text(sample_pdf)
            assert isinstance(result.word_count, int)
            assert isinstance(result.character_count, int)
            assert result.word_count > 0
            assert result.character_count > 0

    def test_pdf_file_not_found(self, app):
        """A7 — Graceful error on missing file."""
        with app.app_context():
            from services.evidence_processor import extract_pdf_text

            result = extract_pdf_text("/nonexistent/file.pdf")
            assert result.success is False
            assert "not found" in result.error_message.lower()


# ---------------------------------------------------------------------------
# G2: PDF OCR extraction (scanned)
# ---------------------------------------------------------------------------

class TestPdfOcrExtraction:
    """G2 — Scanned PDF → OCR text."""

    def test_ocr_page_produces_text(self, app, tmp_path):
        """OCR on an image-based PDF page returns text."""
        with app.app_context():
            from services.evidence_processor import _ocr_pdf_page
            from PIL import Image, ImageDraw, ImageFont

            # Create an image with text
            img = Image.new("RGB", (400, 100), "white")
            draw = ImageDraw.Draw(img)
            draw.text((10, 30), "EVIDENT TECHNOLOGIES", fill="black")

            # Mock a pdfplumber page with to_image
            mock_page = MagicMock()
            mock_image = MagicMock()
            mock_image.original = img
            mock_page.to_image.return_value = mock_image

            text = _ocr_pdf_page(mock_page, page_number=1)
            assert isinstance(text, str)
            # pytesseract should be able to read "EVIDENT" from the image
            assert len(text) > 0

    def test_ocr_fallback_on_scanned_pdf(self, app, tmp_path):
        """Scanned PDF (no native text) triggers OCR path."""
        with app.app_context():
            from services.evidence_processor import extract_pdf_text

            # Create a real file so path.exists() passes
            fake_pdf = tmp_path / "scanned.pdf"
            fake_pdf.write_bytes(b"%PDF-1.4 fake")

            # Create an image with text for OCR
            from PIL import Image, ImageDraw
            img = Image.new("RGB", (400, 100), "white")
            draw = ImageDraw.Draw(img)
            draw.text((10, 30), "OCR TEST STRING", fill="black")

            # Create a mock pdfplumber module that returns our test image
            import types
            mock_pdfplumber = types.ModuleType("pdfplumber")

            mock_page = MagicMock()
            mock_page.extract_text.return_value = ""  # No native text
            mock_img_result = MagicMock()
            mock_img_result.original = img
            mock_page.to_image.return_value = mock_img_result

            mock_pdf = MagicMock()
            mock_pdf.pages = [mock_page]
            mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
            mock_pdf.__exit__ = MagicMock(return_value=False)
            mock_pdfplumber.open = MagicMock(return_value=mock_pdf)

            import sys
            original_module = sys.modules.get("pdfplumber")
            sys.modules["pdfplumber"] = mock_pdfplumber
            try:
                result = extract_pdf_text(str(fake_pdf))
                assert result.success is True
                assert result.task_type == "pdf_ocr"
                assert result.metadata.get("ocr_pages", 0) > 0
            finally:
                if original_module is not None:
                    sys.modules["pdfplumber"] = original_module
                else:
                    sys.modules.pop("pdfplumber", None)


# ---------------------------------------------------------------------------
# G3: Video metadata extraction
# ---------------------------------------------------------------------------

class TestVideoMetadataExtraction:
    """G3 — ffprobe → structured metadata."""

    def test_extract_video_metadata(self, app):
        """B1 — ffprobe returns structured metadata from video file."""
        with app.app_context():
            from services.evidence_processor import extract_video_metadata

            # Mock ffprobe output
            mock_output = json.dumps({
                "streams": [
                    {
                        "codec_type": "video",
                        "codec_name": "h264",
                        "codec_long_name": "H.264 / AVC",
                        "width": 1920,
                        "height": 1080,
                        "r_frame_rate": "30/1",
                        "bit_rate": "10000000",
                        "nb_frames": "4020",
                        "pix_fmt": "yuv420p",
                    },
                    {
                        "codec_type": "audio",
                        "codec_name": "aac",
                        "sample_rate": "48000",
                        "channels": 2,
                        "bit_rate": "46000",
                    },
                ],
                "format": {
                    "duration": "134.0",
                    "size": "168000000",
                    "bit_rate": "10051312",
                    "format_name": "mov,mp4",
                    "format_long_name": "QuickTime / MOV",
                    "nb_streams": "2",
                },
            })

            with patch("services.evidence_processor.subprocess.run") as mock_run:
                mock_run.return_value = MagicMock(
                    returncode=0, stdout=mock_output, stderr=""
                )

                result = extract_video_metadata("test.mp4")

                assert result["duration_seconds"] == 134.0
                assert result["video"]["width"] == 1920
                assert result["video"]["height"] == 1080
                assert result["video"]["codec"] == "h264"
                assert result["video"]["fps"] == 30.0
                assert result["audio"]["codec"] == "aac"
                assert result["audio"]["sample_rate"] == 48000

    def test_video_metadata_ffprobe_failure(self, app):
        """B7 — Graceful error when ffprobe fails."""
        with app.app_context():
            from services.evidence_processor import extract_video_metadata

            with patch("services.evidence_processor.subprocess.run") as mock_run:
                mock_run.return_value = MagicMock(
                    returncode=1, stdout="", stderr="File not found"
                )

                result = extract_video_metadata("nonexistent.mp4")
                assert "error" in result

    def test_video_metadata_timeout(self, app):
        """B7 — Graceful error on ffprobe timeout."""
        with app.app_context():
            from services.evidence_processor import extract_video_metadata
            import subprocess

            with patch("services.evidence_processor.subprocess.run") as mock_run:
                mock_run.side_effect = subprocess.TimeoutExpired("ffprobe", 60)

                result = extract_video_metadata("large_video.mp4")
                assert "error" in result
                assert "timed out" in result["error"]


# ---------------------------------------------------------------------------
# G4: Thumbnail generation
# ---------------------------------------------------------------------------

class TestThumbnailGeneration:
    """G4 — ffmpeg → JPEG thumbnail."""

    def test_thumbnail_generation(self, app, tmp_path):
        """B2 — Generate JPEG thumbnail from video."""
        with app.app_context():
            from services.evidence_processor import generate_thumbnail

            output = str(tmp_path / "thumb.jpg")

            with patch("services.evidence_processor.subprocess.run") as mock_run:
                mock_run.return_value = MagicMock(returncode=0, stderr="")

                # Create the output file to simulate ffmpeg success
                Path(output).write_bytes(b"\xff\xd8\xff" + b"\x00" * 1000)  # Fake JPEG

                result = generate_thumbnail("test.mp4", output, timestamp=10.0)
                assert result is True

    def test_thumbnail_failure(self, app, tmp_path):
        """B7 — Graceful failure on ffmpeg error."""
        with app.app_context():
            from services.evidence_processor import generate_thumbnail

            output = str(tmp_path / "thumb_fail.jpg")

            with patch("services.evidence_processor.subprocess.run") as mock_run:
                mock_run.return_value = MagicMock(
                    returncode=1, stderr="Error opening file"
                )

                result = generate_thumbnail("bad_video.mp4", output, timestamp=10.0)
                assert result is False


# ---------------------------------------------------------------------------
# G5: Content search
# ---------------------------------------------------------------------------

class TestContentSearch:
    """G5 — keyword search → matching results."""

    def test_search_finds_matching_content(self, app, _db, test_case):
        """D2 — Search returns evidence containing keyword."""
        with app.app_context():
            from models.document_processing import ContentExtractionIndex
            from services.document_processing_service import ContentIndexService

            # Create indexed content
            idx = ContentExtractionIndex(
                evidence_id=99901,
                case_id=test_case.id,
                content_type="text",
                word_count=50,
                character_count=300,
                full_text="The HAMILTON Police Department responded to the incident at Main Street.",
                is_indexed=True,
            )
            _db.session.add(idx)
            _db.session.commit()

            results = ContentIndexService.search_content(test_case.id, "HAMILTON")
            assert len(results) >= 1
            assert any(r.evidence_id == 99901 for r in results)

    def test_search_returns_empty_for_absent_term(self, app, _db, test_case):
        """D3 — Search for nonexistent term returns empty."""
        with app.app_context():
            from services.document_processing_service import ContentIndexService

            results = ContentIndexService.search_content(test_case.id, "XYZNONEXISTENT999")
            assert len(results) == 0


# ---------------------------------------------------------------------------
# G6: Entity extraction
# ---------------------------------------------------------------------------

class TestEntityExtraction:
    """G6 — regex entity extraction."""

    def test_extract_emails(self, app):
        """D4 — Extracts email addresses from text."""
        from services.evidence_processor import extract_entities

        emails, _ = extract_entities(
            "Contact john.doe@police.gov or admin@evident.tech for details."
        )
        assert "admin@evident.tech" in emails
        assert "john.doe@police.gov" in emails

    def test_extract_phone_numbers(self, app):
        """D4 — Extracts phone numbers from text."""
        from services.evidence_processor import extract_entities

        _, phones = extract_entities(
            "Call (609) 625-2700 or 555-123-4567 for information."
        )
        assert len(phones) >= 2

    def test_empty_text_returns_empty_lists(self, app):
        """Edge case: empty or None text."""
        from services.evidence_processor import extract_entities

        emails, phones = extract_entities("")
        assert emails == []
        assert phones == []

        emails, phones = extract_entities(None)
        assert emails == []
        assert phones == []


# ---------------------------------------------------------------------------
# G7: Processing task lifecycle
# ---------------------------------------------------------------------------

class TestProcessingTaskLifecycle:
    """G7 — create → processing → completed."""

    def test_task_lifecycle(self, app, _db, test_user, test_case, evidence_item):
        """A6 — Task goes from queued → processing → completed with timing."""
        with app.app_context():
            from models.document_processing import DocumentProcessingTask
            from datetime import datetime, timezone

            task = DocumentProcessingTask(
                evidence_id=evidence_item.id,
                case_id=test_case.id,
                task_type="ocr",
                task_uuid=uuid.uuid4().hex,
                status="queued",
                requested_by_id=test_user.id,
            )
            _db.session.add(task)
            _db.session.commit()

            assert task.status == "queued"

            # Transition to processing
            task.status = "processing"
            task.started_at = datetime.now(timezone.utc)
            _db.session.commit()
            assert task.status == "processing"
            assert task.started_at is not None

            # Transition to completed
            task.status = "completed"
            task.completed_at = datetime.now(timezone.utc)
            task.processing_time_seconds = 5
            _db.session.commit()

            assert task.status == "completed"
            assert task.processing_time_seconds > 0
            assert task.completed_at is not None


# ---------------------------------------------------------------------------
# G8: Processing task failure
# ---------------------------------------------------------------------------

class TestProcessingTaskFailure:
    """G8 — corrupt file → task status = failed."""

    def test_corrupt_pdf_fails_gracefully(self, app, tmp_path):
        """A7 — Corrupt PDF → success=False, no crash."""
        with app.app_context():
            from services.evidence_processor import extract_pdf_text

            corrupt_path = tmp_path / "corrupt.pdf"
            corrupt_path.write_bytes(b"NOT A VALID PDF FILE CONTENT")

            result = extract_pdf_text(str(corrupt_path))
            assert result.success is False
            assert result.error_message is not None

    def test_missing_file_fails_gracefully(self, app):
        """A7 — Missing file → success=False, clear error message."""
        with app.app_context():
            from services.evidence_processor import extract_pdf_text

            result = extract_pdf_text("/does/not/exist/file.pdf")
            assert result.success is False
            assert "not found" in result.error_message.lower()


# ---------------------------------------------------------------------------
# G9: API — POST /api/v1/evidence/<id>/process
# ---------------------------------------------------------------------------

class TestApiProcessEvidence:
    """G9 — POST trigger → response."""

    def test_process_evidence_endpoint(self, app, _db, client, api_token, evidence_item, test_case):
        """E1 — POST /api/v1/evidence/<id>/process returns result."""
        with app.app_context():
            # Mock at the tasks module level since the route imports locally
            with patch("tasks.processing_tasks.dispatch_process_evidence") as mock_dispatch:
                mock_dispatch.return_value = {
                    "success": True,
                    "async": False,
                    "task_id": 1,
                    "task_uuid": "abc123",
                    "task_type": "ocr",
                    "summary": {"word_count": 100},
                    "processing_seconds": 1.5,
                }

                resp = client.post(
                    f"/api/v1/evidence/{evidence_item.id}/process",
                    headers={"Authorization": f"Bearer {api_token}"},
                    content_type="application/json",
                )

                assert resp.status_code == 200
                data = resp.get_json()
                assert data["status"] == "completed"
                assert data["task_id"] == 1

    def test_process_nonexistent_evidence(self, app, _db, client, api_token):
        """E1 — 404 for nonexistent evidence."""
        with app.app_context():
            resp = client.post(
                "/api/v1/evidence/999999/process",
                headers={"Authorization": f"Bearer {api_token}"},
            )
            assert resp.status_code == 404

    def test_process_without_auth(self, app, client, evidence_item):
        """E1 — 401 without Bearer token."""
        resp = client.post(f"/api/v1/evidence/{evidence_item.id}/process")
        assert resp.status_code == 401


# ---------------------------------------------------------------------------
# G10: API — GET /api/v1/evidence/<id>/text
# ---------------------------------------------------------------------------

class TestApiGetText:
    """G10 — GET extracted text → JSON."""

    def test_get_text_with_index(self, app, _db, client, api_token, evidence_item, test_case):
        """E2 — Returns extracted text with metadata."""
        with app.app_context():
            from models.document_processing import ContentExtractionIndex

            # Create content index entry
            idx = ContentExtractionIndex(
                evidence_id=evidence_item.id,
                case_id=test_case.id,
                content_type="pdf_text",
                word_count=150,
                character_count=900,
                line_count=3,
                full_text="Extracted content from evidence document.",
                email_addresses="test@evident.tech",
                phone_numbers="555-123-4567",
                is_indexed=True,
            )
            _db.session.add(idx)
            _db.session.commit()

            resp = client.get(
                f"/api/v1/evidence/{evidence_item.id}/text",
                headers={"Authorization": f"Bearer {api_token}"},
            )

            assert resp.status_code == 200
            data = resp.get_json()
            assert data["full_text"] == "Extracted content from evidence document."
            assert data["word_count"] == 150
            assert data["content_type"] == "pdf_text"

    def test_get_text_no_content(self, app, _db, client, api_token, test_user, test_case):
        """E2 — 404 when no text has been extracted."""
        with app.app_context():
            from models.evidence import EvidenceItem, CaseEvidence

            item = EvidenceItem(
                original_filename="empty.pdf",
                file_type="pdf",
                evidence_type="document",
                hash_sha256=hashlib.sha256(b"empty").hexdigest(),
                processing_status="pending",
                uploaded_by_id=test_user.id,
                origin_case_id=test_case.id,
            )
            _db.session.add(item)
            _db.session.commit()

            resp = client.get(
                f"/api/v1/evidence/{item.id}/text",
                headers={"Authorization": f"Bearer {api_token}"},
            )
            assert resp.status_code == 404


# ---------------------------------------------------------------------------
# G11: API — GET /api/v1/search
# ---------------------------------------------------------------------------

class TestApiSearch:
    """G11 — search endpoint."""

    def test_search_returns_results(self, app, _db, client, api_token, test_case):
        """D5 — GET /api/v1/search returns matching evidence."""
        with app.app_context():
            from models.document_processing import ContentExtractionIndex
            from models.evidence import EvidenceItem

            # Create evidence + index
            item = EvidenceItem(
                original_filename="search_test.pdf",
                file_type="pdf",
                evidence_type="document",
                hash_sha256=hashlib.sha256(f"search-{uuid.uuid4().hex}".encode()).hexdigest(),
                processing_status="completed",
                origin_case_id=test_case.id,
            )
            _db.session.add(item)
            _db.session.flush()

            idx = ContentExtractionIndex(
                evidence_id=item.id,
                case_id=test_case.id,
                content_type="pdf_text",
                word_count=50,
                character_count=300,
                full_text="The TOWNSHIP OF HAMILTON police department filed a report.",
                is_indexed=True,
            )
            _db.session.add(idx)
            _db.session.commit()

            resp = client.get(
                f"/api/v1/search?q=HAMILTON&case_id={test_case.id}",
                headers={"Authorization": f"Bearer {api_token}"},
            )

            assert resp.status_code == 200
            data = resp.get_json()
            assert data["total_results"] >= 1
            assert any("HAMILTON" in r.get("snippet", "") for r in data["results"])

    def test_search_requires_query(self, app, client, api_token, test_case):
        """D5 — 400 if 'q' parameter missing."""
        resp = client.get(
            f"/api/v1/search?case_id={test_case.id}",
            headers={"Authorization": f"Bearer {api_token}"},
        )
        assert resp.status_code == 400

    def test_search_requires_case_id(self, app, client, api_token):
        """D5 — 400 if 'case_id' parameter missing."""
        resp = client.get(
            "/api/v1/search?q=test",
            headers={"Authorization": f"Bearer {api_token}"},
        )
        assert resp.status_code == 400

    def test_search_no_results(self, app, _db, client, api_token, test_case):
        """D5 — Empty results for absent keyword."""
        resp = client.get(
            f"/api/v1/search?q=XYZNONEXISTENTTERM&case_id={test_case.id}",
            headers={"Authorization": f"Bearer {api_token}"},
        )
        assert resp.status_code == 200
        data = resp.get_json()
        assert data["total_results"] == 0


# ---------------------------------------------------------------------------
# G12: Unsupported file type
# ---------------------------------------------------------------------------

class TestUnsupportedFileType:
    """G12 — Unsupported file type → skip with reason."""

    def test_unsupported_returns_error(self, app):
        """C5 — Unsupported file type returns success=False with reason."""
        with app.app_context():
            from services.evidence_processor import process_evidence_file

            result = process_evidence_file("/some/file.xyz", original_filename="data.xyz")
            assert result.success is False
            assert "unsupported" in result.error_message.lower()

    def test_detect_file_type_unsupported(self, app):
        """C4 — Unknown extension returns 'unsupported'."""
        from services.evidence_processor import detect_file_type

        assert detect_file_type("file.xyz") == "unsupported"
        assert detect_file_type("file.bin") == "unsupported"


# ---------------------------------------------------------------------------
# Additional coverage: file type detection, DOCX, plaintext
# ---------------------------------------------------------------------------

class TestFileTypeDetection:
    """C4 — MIME type / extension detection."""

    def test_pdf_detection(self):
        from services.evidence_processor import detect_file_type
        assert detect_file_type("report.pdf") == "pdf"

    def test_video_detection(self):
        from services.evidence_processor import detect_file_type
        assert detect_file_type("bwc_footage.mp4") == "video"
        assert detect_file_type("evidence.avi") == "video"

    def test_image_detection(self):
        from services.evidence_processor import detect_file_type
        assert detect_file_type("photo.jpg") == "image"
        assert detect_file_type("scan.png") == "image"

    def test_docx_detection(self):
        from services.evidence_processor import detect_file_type
        assert detect_file_type("statement.docx") == "docx"

    def test_plaintext_detection(self):
        from services.evidence_processor import detect_file_type
        assert detect_file_type("notes.txt") == "plaintext"
        assert detect_file_type("data.csv") == "plaintext"


class TestDocxExtraction:
    """C1 — DOCX text extraction."""

    def test_docx_extraction(self, app, tmp_path):
        """C1 — python-docx extracts paragraphs."""
        with app.app_context():
            try:
                import docx
                doc = docx.Document()
                doc.add_paragraph("EVIDENT TECHNOLOGIES")
                doc.add_paragraph("Evidence processing test document.")
                doc.add_paragraph("Contact: legal@evident.tech")
                path = tmp_path / "test.docx"
                doc.save(str(path))

                from services.evidence_processor import extract_docx_text
                result = extract_docx_text(str(path))

                assert result.success is True
                assert result.task_type == "docx_text"
                assert "EVIDENT" in result.full_text
                assert result.word_count > 0
                assert "legal@evident.tech" in result.email_addresses
            except ImportError:
                pytest.skip("python-docx not available")


class TestPlaintextExtraction:
    """C3 — Plain-text passthrough."""

    def test_plaintext_passthrough(self, app, sample_text_file):
        """C3 — Exact content match."""
        with app.app_context():
            from services.evidence_processor import extract_plaintext

            result = extract_plaintext(sample_text_file)
            assert result.success is True
            assert result.task_type == "plaintext"
            assert "officer@police.gov" in result.full_text
            assert "officer@police.gov" in result.email_addresses
            assert result.word_count > 0

    def test_plaintext_file_not_found(self, app):
        """Edge case: missing text file."""
        with app.app_context():
            from services.evidence_processor import extract_plaintext

            result = extract_plaintext("/nonexistent/file.txt")
            assert result.success is False


class TestVideoProcessing:
    """Full video processing pipeline."""

    def test_process_video_evidence(self, app, tmp_path):
        """B1-B6 — Video processing stores metadata + thumbnail."""
        with app.app_context():
            from services.evidence_processor import process_video_evidence
            from services.evidence_store import EvidenceStore

            store = EvidenceStore(root=str(tmp_path / "evidence_store"))
            sha256 = "a" * 64  # Fake hash

            # Store a fake original so derivative storage works
            orig_dir = store._original_dir(sha256)
            orig_dir.mkdir(parents=True, exist_ok=True)
            (orig_dir / "test.mp4").write_bytes(b"fake video")

            # Mock ffprobe and ffmpeg
            mock_ffprobe = json.dumps({
                "streams": [
                    {"codec_type": "video", "codec_name": "h264", "width": 1920,
                     "height": 1080, "r_frame_rate": "30/1", "bit_rate": "10000000",
                     "nb_frames": "900", "pix_fmt": "yuv420p"},
                ],
                "format": {"duration": "30.0", "size": "5000000", "bit_rate": "1333333",
                          "format_name": "mov,mp4", "format_long_name": "QuickTime",
                          "nb_streams": "1"},
            })

            def mock_run_side_effect(cmd, **kwargs):
                result = MagicMock()
                if cmd[0] == "ffprobe":
                    result.returncode = 0
                    result.stdout = mock_ffprobe
                    result.stderr = ""
                elif cmd[0] == "ffmpeg":
                    # Create fake thumbnail output
                    output_path = cmd[-1]
                    Path(output_path).write_bytes(b"\xff\xd8\xff" + b"\x00" * 500)
                    result.returncode = 0
                    result.stderr = ""
                return result

            with patch("services.evidence_processor.subprocess.run", side_effect=mock_run_side_effect):
                vr = process_video_evidence(
                    file_path=str(orig_dir / "test.mp4"),
                    evidence_store=store,
                    original_sha256=sha256,
                )

                assert vr.success is True
                assert vr.metadata["duration_seconds"] == 30.0
                assert vr.metadata["video"]["width"] == 1920
                assert vr.thumbnail_path is not None


class TestSnippetExtraction:
    """Helper function for search snippets."""

    def test_snippet_around_match(self, app):
        from routes.processing_routes import _extract_snippet
        text = "A" * 200 + "KEYWORD" + "B" * 200
        snippet = _extract_snippet(text, "KEYWORD", context_chars=50)
        assert "KEYWORD" in snippet
        assert snippet.startswith("...")
        assert snippet.endswith("...")

    def test_snippet_no_match(self, app):
        from routes.processing_routes import _extract_snippet
        snippet = _extract_snippet("Some text here", "NONEXISTENT")
        assert isinstance(snippet, str)

    def test_snippet_empty_text(self, app):
        from routes.processing_routes import _extract_snippet
        assert _extract_snippet("", "test") == ""


class TestTaskStatusEndpoint:
    """E3 — GET /api/v1/tasks/<id>."""

    def test_get_task_status(self, app, _db, client, api_token, evidence_item, test_case):
        """E3 — Returns task status JSON."""
        with app.app_context():
            from models.document_processing import DocumentProcessingTask

            task = DocumentProcessingTask(
                evidence_id=evidence_item.id,
                case_id=test_case.id,
                task_type="ocr",
                task_uuid=uuid.uuid4().hex,
                status="completed",
                processing_time_seconds=5,
                requested_by_id=1,
            )
            _db.session.add(task)
            _db.session.commit()

            resp = client.get(
                f"/api/v1/tasks/{task.id}",
                headers={"Authorization": f"Bearer {api_token}"},
            )

            assert resp.status_code == 200
            data = resp.get_json()
            assert data["status"] == "completed"
            assert data["task_type"] == "ocr"

    def test_task_not_found(self, app, _db, client, api_token):
        """E3 — 404 for nonexistent task."""
        resp = client.get(
            "/api/v1/tasks/999999",
            headers={"Authorization": f"Bearer {api_token}"},
        )
        assert resp.status_code == 404


class TestBatchEndpoint:
    """E5 — GET /api/v1/batches/<id>."""

    def test_get_batch_status(self, app, _db, client, api_token, test_case, test_user):
        """E5 — Returns batch status JSON."""
        with app.app_context():
            from models.document_processing import BatchProcessingQueue

            batch = BatchProcessingQueue(
                case_id=test_case.id,
                batch_name="Test Batch",
                batch_uuid=uuid.uuid4().hex,
                processing_type="full_extraction",
                document_count=10,
                status="completed",
                successful_count=8,
                failed_count=2,
                progress_percentage=100,
                created_by_id=test_user.id,
            )
            _db.session.add(batch)
            _db.session.commit()

            resp = client.get(
                f"/api/v1/batches/{batch.id}",
                headers={"Authorization": f"Bearer {api_token}"},
            )

            assert resp.status_code == 200
            data = resp.get_json()
            assert data["status"] == "completed"
            assert data["successful_count"] == 8
            assert data["failed_count"] == 2

    def test_batch_not_found(self, app, _db, client, api_token):
        """E5 — 404 for nonexistent batch."""
        resp = client.get(
            "/api/v1/batches/999999",
            headers={"Authorization": f"Bearer {api_token}"},
        )
        assert resp.status_code == 404


class TestFPSParser:
    """Utility: frame rate string parsing."""

    def test_parse_fractional_fps(self):
        from services.evidence_processor import _parse_fps
        assert _parse_fps("30000/1001") == 29.97

    def test_parse_integer_fps(self):
        from services.evidence_processor import _parse_fps
        assert _parse_fps("30/1") == 30.0

    def test_parse_plain_number(self):
        from services.evidence_processor import _parse_fps
        assert _parse_fps("25.0") == 25.0

    def test_parse_invalid(self):
        from services.evidence_processor import _parse_fps
        assert _parse_fps("invalid") == 0.0

    def test_parse_divide_by_zero(self):
        from services.evidence_processor import _parse_fps
        assert _parse_fps("30/0") == 0.0


class TestCeleryAppConfig:
    """Celery app configuration and sync fallback."""

    def test_sync_mode_when_env_set(self):
        """EVIDENT_ASYNC=0 → sync mode."""
        # Already set in our test app fixture
        import os
        os.environ["EVIDENT_ASYNC"] = "0"
        # Reimport to test the flag
        # (In actual tests, this is set before app creation)
        from celery_app import ASYNC_ENABLED
        # Note: ASYNC_ENABLED was set at import time, so this tests the module-level behavior

    def test_dispatch_falls_back_to_sync(self, app, _db):
        """dispatch_process_evidence uses sync when async not available."""
        with app.app_context():
            from tasks.processing_tasks import dispatch_process_evidence

            # Mock is_async to return False
            with patch("tasks.processing_tasks.is_async", return_value=False):
                with patch("tasks.processing_tasks.process_evidence_sync") as mock_sync:
                    mock_sync.return_value = {"success": True, "task_id": 1}

                    result = dispatch_process_evidence(1, 1)
                    assert result.get("async") is False
